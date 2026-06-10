from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

titulo = doc.add_heading('APRENDIENDO PYTHON Y ANGULAR', 0)
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitulo = doc.add_heading('Basado en el Proyecto FoodGestor', level=2)
subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

parrafo = doc.add_paragraph('Guia Practica de Desarrollo Full Stack')
parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
parrafo.runs[0].italic = True

doc.add_paragraph()
fecha = doc.add_paragraph('Junio 2026')
fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

doc.add_heading('1. Introduccion', level=1)
doc.add_paragraph('Este documento te ensena como aprender Python y Angular de manera practica, usando FoodGestor como ejemplo real.')

doc.add_heading('Componentes Principales:', level=2)
doc.add_paragraph('Backend en Python con Flask (API REST)', style='List Bullet')
doc.add_paragraph('Frontend en Angular 21 (Aplicacion Web)', style='List Bullet')
doc.add_paragraph('Base de datos PostgreSQL en la nube', style='List Bullet')
doc.add_paragraph('Despliegue como PWA (Progressive Web App)', style='List Bullet')

doc.add_page_break()
doc.add_heading('2. Python y Flask: El Backend', level=1)
doc.add_paragraph('Python es un lenguaje legible y poderoso. Flask es un framework minimalista perfecto para aprender.')

doc.add_heading('Estructura del Proyecto', level=2)
doc.add_paragraph('app/__init__.py: Configuracion de Flask', style='List Bullet')
doc.add_paragraph('app/models/: Definicion de datos', style='List Bullet')
doc.add_paragraph('app/routes/: Endpoints del API', style='List Bullet')

doc.add_heading('Modelos: Estructura de Datos', level=2)
doc.add_paragraph('Usuario: email, password_hash, nombre, limites', style='List Bullet')
doc.add_paragraph('Alimento: nombre, calorias, proteinas, grasas, usuario_id', style='List Bullet')
doc.add_paragraph('Racion: cantidad, alimento_id, usuario_id, fecha', style='List Bullet')

doc.add_heading('Rutas (Endpoints)', level=2)
doc.add_paragraph('POST /api/auth/login - Usuario inicia sesion', style='List Bullet')
doc.add_paragraph('GET /api/alimentos - Obtener alimentos del usuario', style='List Bullet')
doc.add_paragraph('POST /api/raciones - Crear una nueva racion', style='List Bullet')

doc.add_heading('Autenticacion JWT', level=2)
doc.add_paragraph('Usuario se registra con email y password', style='List Number')
doc.add_paragraph('Backend crea un token JWT', style='List Number')
doc.add_paragraph('Frontend guarda el token en localStorage', style='List Number')
doc.add_paragraph('Cada solicitud incluye el token', style='List Number')

doc.add_page_break()
doc.add_heading('3. Angular: El Frontend', level=1)
doc.add_paragraph('Angular proporciona una estructura completa para aplicaciones web.')

doc.add_heading('Estructura del Proyecto', level=2)
doc.add_paragraph('components/: Componentes UI', style='List Bullet')
doc.add_paragraph('services/: Servicios de comunicacion', style='List Bullet')
doc.add_paragraph('guards/: Proteccion de rutas', style='List Bullet')
doc.add_paragraph('interceptors/: Middleware HTTP', style='List Bullet')

doc.add_heading('Componentes', level=2)
doc.add_paragraph('Un componente consta de:')
doc.add_paragraph('.ts: Logica (TypeScript)', style='List Bullet')
doc.add_paragraph('.html: Estructura (HTML)', style='List Bullet')
doc.add_paragraph('.css: Estilos (CSS)', style='List Bullet')

doc.add_heading('Servicios', level=2)
doc.add_paragraph('Los servicios manejan la comunicacion con el backend. El AuthService se encarga de login, registro y perfil.')

doc.add_page_break()
doc.add_heading('4. Integracion Frontend-Backend', level=1)

doc.add_heading('Flujo: Login', level=2)
doc.add_paragraph('Usuario escribe email y password', style='List Number')
doc.add_paragraph('Frontend llama a AuthService.login()', style='List Number')
doc.add_paragraph('AuthService hace POST /api/auth/login', style='List Number')
doc.add_paragraph('Backend valida y crea JWT token', style='List Number')
doc.add_paragraph('Frontend guarda token en localStorage', style='List Number')
doc.add_paragraph('Frontend redirige a /perfil', style='List Number')

doc.add_heading('Flujo: Obtener Alimentos', level=2)
doc.add_paragraph('Frontend llama a AlimentosService.obtener()', style='List Number')
doc.add_paragraph('AuthInterceptor agrega JWT al header', style='List Number')
doc.add_paragraph('Backend verifica token y obtiene usuario_id', style='List Number')
doc.add_paragraph('Backend consulta BD por usuario_id', style='List Number')
doc.add_paragraph('Frontend recibe JSON y lo muestra', style='List Number')

doc.add_page_break()
doc.add_heading('5. Mejores Practicas', level=1)

doc.add_heading('Backend (Python)', level=2)
doc.add_paragraph('Usa blueprints para organizar rutas', style='List Bullet')
doc.add_paragraph('Valida entrada del usuario siempre', style='List Bullet')
doc.add_paragraph('Usa variables de entorno para secretos', style='List Bullet')
doc.add_paragraph('Retorna codigos HTTP correctos', style='List Bullet')

doc.add_heading('Frontend (Angular)', level=2)
doc.add_paragraph('Componentes pequenos (una tarea)', style='List Bullet')
doc.add_paragraph('Logica en servicios, no en componentes', style='List Bullet')
doc.add_paragraph('Usa guards para proteger rutas', style='List Bullet')
doc.add_paragraph('Evita usar "any" - usa TypeScript', style='List Bullet')

doc.add_heading('General', level=2)
doc.add_paragraph('Usa git para control de versiones', style='List Bullet')
doc.add_paragraph('Escribe commits descriptivos', style='List Bullet')
doc.add_paragraph('Documenta tu codigo', style='List Bullet')
doc.add_paragraph('Prueba localmente antes de desplegar', style='List Bullet')

doc.add_page_break()
doc.add_heading('Conclusion', level=1)
doc.add_paragraph('Ahora entiendes como funcionan Python, Flask y Angular juntos.')
doc.add_paragraph('Como crear un API REST en Python', style='List Bullet')
doc.add_paragraph('Como construir una UI en Angular', style='List Bullet')
doc.add_paragraph('Como integrar ambos con JWT', style='List Bullet')
doc.add_paragraph('Como desplegar en produccion', style='List Bullet')

final = doc.add_paragraph()
final.add_run('Felicitaciones! Ya puedes crear aplicaciones web completas.')
final.alignment = WD_ALIGN_PARAGRAPH.CENTER
final.runs[0].bold = True
final.runs[0].italic = True

output_path = 'C:\\Users\\Joza\\Documents\\Proyectos\\FoodGestor\\APRENDIENDO_PYTHON_Y_ANGULAR.docx'
doc.save(output_path)
print("Documento creado exitosamente!")
